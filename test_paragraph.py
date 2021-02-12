import unittest
import os
import docx
import paragraph

class TestParagraph(unittest.TestCase):
    TESTDATA_FILENAME = os.path.join(os.path.dirname(__file__), 'test.docx')
    TESTDATA_OUT_FILENAME = os.path.join(os.path.dirname(__file__), 'out.docx')

    def test_runs_returns_runs_with_hyperlinks(self):
        doc = docx.Document(self.TESTDATA_FILENAME)
        actuals = []
        for p in doc.paragraphs:
            for r in p.runs:
                actuals.append(r.text)

        expecteds = ['みかん', 'りんご', 'アセロラ', 'アケビ', 'あんず', 'アボカド']
        self.assertListEqual(expecteds, actuals)

    def test_runs_returns_runs_with_hyperlinks(self):
        repl = dict()
        repl['みかん'] = 'orange'
        repl['りんご'] = 'apple'
        repl['アセロラ'] = 'acerola'
        repl['アケビ'] = 'akebia'
        repl['あんず'] = 'apricot'
        repl['アボカド'] = 'avocado'

        doc = docx.Document(self.TESTDATA_FILENAME)
        for p in doc.paragraphs:
            for r in p.runs:
                for src, dst in repl.items():
                    r.text = r.text.replace(src, dst)
        doc.save(self.TESTDATA_OUT_FILENAME)

        doc = docx.Document(self.TESTDATA_OUT_FILENAME)
        actuals = []
        for p in doc.paragraphs:
            for r in p.runs:
                actuals.append(r.text)

        expecteds = ['orange', 'apple', 'acerola', 'akebia', 'apricot', 'avocado']
        self.assertListEqual(expecteds, actuals)

    def tearDown(self):
        if os.path.exists(self.TESTDATA_OUT_FILENAME):
            os.remove(self.TESTDATA_OUT_FILENAME)

if __name__ == "__main__":
    unittest.main()